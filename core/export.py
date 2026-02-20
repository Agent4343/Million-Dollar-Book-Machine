"""
Book Export Module

Generates professional book formats:
- Word Document (.docx)
- EPUB for Kindle/eReaders
- Markdown
"""

import io
import re
from typing import Dict, Any, List, Optional
from datetime import datetime, timezone


def _get_best_chapters(project, chapters_override: Optional[List[Dict[str, Any]]] = None) -> List[Dict[str, Any]]:
    if isinstance(chapters_override, list):
        return chapters_override
    chapters = project.manuscript.get("chapters", [])
    if isinstance(chapters, list):
        return chapters
    return []


def _front_matter_defaults(project) -> Dict[str, Any]:
    c = project.user_constraints or {}
    year = c.get("copyright_year") or datetime.now(timezone.utc).year
    return {
        "author_name": c.get("author_name") or c.get("pen_name") or "Author Name",
        "publisher_name": c.get("publisher_name") or "",
        "copyright_year": year,
        "include_disclaimer": bool(c.get("include_disclaimer", True)),
        "disclaimer_text": c.get("disclaimer_text") or "This is a work of fiction. Names, characters, businesses, places, events, and incidents are either the products of the author’s imagination or used in a fictitious manner.",
        "isbn": c.get("isbn") or "",
        "rights_statement": c.get("rights_statement") or "All rights reserved.",
    }


def _supplemental_matter(project) -> Dict[str, Any]:
    c = project.user_constraints or {}
    also_by = c.get("also_by") or c.get("also_by_titles") or []
    if isinstance(also_by, str):
        also_by = [t.strip() for t in also_by.split(",") if t.strip()]
    if not isinstance(also_by, list):
        also_by = []
    also_by = [str(t).strip() for t in also_by if str(t).strip()]
    return {
        "also_by": also_by,
        "acknowledgements": (c.get("acknowledgements") or "").strip(),
        "about_author": (c.get("about_author") or c.get("about_author_text") or "").strip(),
        "newsletter_cta": (c.get("newsletter_cta") or "").strip(),
        "newsletter_url": (c.get("newsletter_url") or "").strip(),
    }


def generate_docx(project, include_outline: bool = False, chapters_override: Optional[List[Dict[str, Any]]] = None) -> bytes:
    """
    Generate a Word document from the project manuscript.

    Args:
        project: BookProject instance
        include_outline: Whether to include the development outline

    Returns:
        bytes: The .docx file content
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True

    # Heading 1 for chapters
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True

    # Normal paragraph style
    normal_style = styles['Normal']
    normal_style.font.size = Pt(12)
    normal_style.font.name = 'Times New Roman'

    # === Title Page ===
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(project.title)
    title_run.font.size = Pt(36)
    title_run.font.bold = True

    # Add some space
    for _ in range(3):
        doc.add_paragraph()

    fm = _front_matter_defaults(project)

    # Subtitle/description if available
    if project.user_constraints.get('description'):
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_text = project.user_constraints['description'][:200]
        if len(project.user_constraints['description']) > 200:
            desc_text += '...'
        subtitle.add_run(desc_text).italic = True

    # Genre and word count
    for _ in range(5):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    genre = project.user_constraints.get('genre', 'Fiction').replace('_', ' ').title()
    meta.add_run(f"Genre: {genre}")

    # Page break after title
    doc.add_page_break()

    # === Copyright / Disclaimer Page (KDP recommended) ===
    copyright_heading = doc.add_heading("Copyright", level=1)
    copyright_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    cpara = doc.add_paragraph()
    cpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpara.add_run(f"© {fm['copyright_year']} {fm['author_name']}\n{fm['rights_statement']}")
    if fm.get("isbn"):
        doc.add_paragraph().add_run(f"ISBN: {fm['isbn']}")
    if fm.get("publisher_name"):
        doc.add_paragraph().add_run(f"Publisher: {fm['publisher_name']}")
    if fm.get("include_disclaimer"):
        doc.add_paragraph()
        disc = doc.add_paragraph()
        disc.add_run("Disclaimer: ").bold = True
        doc.add_paragraph(fm.get("disclaimer_text", ""))

    doc.add_page_break()

    sup = _supplemental_matter(project)

    # === Also By (optional) ===
    if sup["also_by"]:
        doc.add_heading("Also By", level=1)
        for t in sup["also_by"]:
            doc.add_paragraph(t)
        doc.add_page_break()

    # Get all agent outputs for use throughout document
    outputs = {}
    for layer in project.layers.values():
        for agent_id, agent_state in layer.agents.items():
            if agent_state.current_output:
                outputs[agent_id] = agent_state.current_output.content

    # === Table of Contents ===
    toc_heading = doc.add_heading('Table of Contents', level=1)

    chapters = _get_best_chapters(project, chapters_override=chapters_override)

    # If no written chapters, use chapter blueprint for TOC
    if not chapters and 'chapter_blueprint' in outputs:
        blueprint = outputs['chapter_blueprint']
        chapter_outline = blueprint.get('chapter_outline', [])
        for ch in chapter_outline:
            ch_num = ch.get('number', '?')
            ch_title = ch.get('title', f'Chapter {ch_num}')
            toc_entry = doc.add_paragraph(f"Chapter {ch_num}: {ch_title}")
            toc_entry.paragraph_format.left_indent = Inches(0.5)
    elif chapters:
        for chapter in sorted(chapters, key=lambda x: x.get('number', 0)):
            ch_num = chapter.get('number', '?')
            ch_title = chapter.get('title', f'Chapter {ch_num}')
            toc_entry = doc.add_paragraph(f"Chapter {ch_num}: {ch_title}")
            toc_entry.paragraph_format.left_indent = Inches(0.5)
    else:
        doc.add_paragraph("No chapters planned yet. Run the pipeline first.")

    doc.add_page_break()

    # === Core Concept Section (always include if available) ===
    if 'concept_definition' in outputs:
        cd = outputs['concept_definition']
        doc.add_heading("Book Concept", level=1)

        if cd.get('one_line_hook'):
            hook_para = doc.add_paragraph()
            hook_para.add_run("Hook: ").bold = True
            hook_para.add_run(cd['one_line_hook'])

        if cd.get('elevator_pitch'):
            doc.add_paragraph()
            pitch_para = doc.add_paragraph()
            pitch_para.add_run("Elevator Pitch: ").bold = True
            doc.add_paragraph(cd['elevator_pitch'])

        if cd.get('core_promise'):
            cp = cd['core_promise']
            doc.add_paragraph()
            doc.add_paragraph(f"Transformation: {cp.get('transformation', 'N/A')}")
            doc.add_paragraph(f"Emotional Payoff: {cp.get('emotional_payoff', 'N/A')}")

        doc.add_page_break()

    # === Character Section (always include if available) ===
    if 'character_architecture' in outputs:
        ca = outputs['character_architecture']
        doc.add_heading("Characters", level=1)

        if ca.get('protagonist_profile'):
            pp = ca['protagonist_profile']
            doc.add_heading("Protagonist", level=2)
            doc.add_paragraph(f"Name: {pp.get('name', 'Unknown')}")
            doc.add_paragraph(f"Role: {pp.get('role', 'N/A')}")
            if pp.get('traits'):
                doc.add_paragraph(f"Traits: {', '.join(pp['traits'])}")
            if pp.get('backstory_wound'):
                doc.add_paragraph(f"Wound: {pp.get('backstory_wound')}")

        if ca.get('protagonist_arc'):
            pa = ca['protagonist_arc']
            doc.add_heading("Character Arc", level=2)
            doc.add_paragraph(f"Starting State: {pa.get('starting_state', 'N/A')}")
            doc.add_paragraph(f"Transformation: {pa.get('transformation', 'N/A')}")
            doc.add_paragraph(f"Ending State: {pa.get('ending_state', 'N/A')}")

        if ca.get('supporting_cast'):
            doc.add_heading("Supporting Cast", level=2)
            for char in ca['supporting_cast'][:5]:
                doc.add_paragraph(f"• {char.get('name', '?')}: {char.get('function', 'N/A')}")

        doc.add_page_break()

    # === Chapters ===
    if chapters:
        doc.add_heading("Manuscript", level=1)
        doc.add_page_break()

        for chapter in sorted(chapters, key=lambda x: x.get('number', 0)):
            ch_num = chapter.get('number', '?')
            ch_title = chapter.get('title', f'Chapter {ch_num}')

            # Chapter heading
            doc.add_heading(f"Chapter {ch_num}: {ch_title}", level=1)

            # Chapter content
            text = chapter.get('text', '')
            if text:
                # Split into paragraphs and add each
                paragraphs = text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # Handle scene breaks
                        if para_text in ['* * *', '---', '***']:
                            scene_break = doc.add_paragraph()
                            scene_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            scene_break.add_run('* * *')
                        else:
                            para = doc.add_paragraph(para_text)
                            para.paragraph_format.first_line_indent = Inches(0.5)
            else:
                doc.add_paragraph("[Chapter content not yet written]").italic = True

            # Page break between chapters
            doc.add_page_break()

    # === Chapter Blueprint (if no written chapters) ===
    elif 'chapter_blueprint' in outputs:
        doc.add_heading("Chapter Outline", level=1)
        doc.add_paragraph("Note: Use the Chapter Writer to generate full prose for each chapter.")
        doc.add_paragraph()

        blueprint = outputs['chapter_blueprint']
        chapter_outline = blueprint.get('chapter_outline', [])

        for ch in chapter_outline:
            ch_num = ch.get('number', '?')
            ch_title = ch.get('title', f'Chapter {ch_num}')

            doc.add_heading(f"Chapter {ch_num}: {ch_title}", level=2)

            if ch.get('chapter_goal'):
                goal_para = doc.add_paragraph()
                goal_para.add_run("Goal: ").bold = True
                goal_para.add_run(ch['chapter_goal'])

            if ch.get('opening_hook'):
                hook_para = doc.add_paragraph()
                hook_para.add_run("Opening: ").bold = True
                hook_para.add_run(ch['opening_hook'])

            if ch.get('closing_hook'):
                close_para = doc.add_paragraph()
                close_para.add_run("Closing: ").bold = True
                close_para.add_run(ch['closing_hook'])

            # Scenes
            if ch.get('scenes'):
                doc.add_paragraph()
                scenes_para = doc.add_paragraph()
                scenes_para.add_run("Scenes:").bold = True
                for scene in ch['scenes']:
                    doc.add_paragraph(
                        f"  {scene.get('scene_number', '?')}. {scene.get('scene_question', 'Scene')} "
                        f"[{scene.get('location', 'Location')}]"
                    )

            doc.add_paragraph()
    else:
        doc.add_heading("No Content Yet", level=1)
        doc.add_paragraph("Run the pipeline to generate your book outline, then use the Chapter Writer to create full chapters.")

    # === Back Matter (optional) ===
    if sup.get("acknowledgements"):
        doc.add_heading("Acknowledgements", level=1)
        doc.add_paragraph(sup["acknowledgements"])
        doc.add_page_break()

    if sup.get("newsletter_cta") or sup.get("newsletter_url"):
        doc.add_heading("Stay in Touch", level=1)
        if sup.get("newsletter_cta"):
            doc.add_paragraph(sup["newsletter_cta"])
        if sup.get("newsletter_url"):
            doc.add_paragraph(sup["newsletter_url"])
        doc.add_page_break()

    if sup.get("about_author"):
        doc.add_heading("About the Author", level=1)
        doc.add_paragraph(sup["about_author"])
        doc.add_page_break()

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_epub(project, chapters_override: Optional[List[Dict[str, Any]]] = None) -> bytes:
    """
    Generate an EPUB file from the project manuscript.
    Compatible with Kindle and other eReaders.

    Args:
        project: BookProject instance

    Returns:
        bytes: The .epub file content
    """
    from ebooklib import epub

    book = epub.EpubBook()

    # Set metadata
    book.set_identifier(f'million-dollar-book-{project.project_id}')
    book.set_title(project.title)
    book.set_language('en')

    # Add author (placeholder - could be made configurable)
    fm = _front_matter_defaults(project)
    book.add_author(fm["author_name"])

    # Add description
    if project.user_constraints.get('description'):
        book.add_metadata('DC', 'description', project.user_constraints['description'])

    # Genre as subject
    genre = project.user_constraints.get('genre', 'Fiction').replace('_', ' ').title()
    book.add_metadata('DC', 'subject', genre)

    # Create chapters
    chapters = _get_best_chapters(project, chapters_override=chapters_override)
    epub_chapters = []
    sup = _supplemental_matter(project)

    # Title page
    title_page = epub.EpubHtml(title='Title Page', file_name='title.xhtml', lang='en')
    title_page.content = f'''
    <html>
    <head><title>{project.title}</title></head>
    <body>
        <div style="text-align: center; margin-top: 30%;">
            <h1 style="font-size: 2.5em;">{project.title}</h1>
            <p style="margin-top: 2em; font-style: italic;">{genre}</p>
        </div>
    </body>
    </html>
    '''
    book.add_item(title_page)
    epub_chapters.append(title_page)

    # Copyright page (KDP recommended)
    copyright_page = epub.EpubHtml(title='Copyright', file_name='copyright.xhtml', lang='en')
    disclaimer_html = ""
    if fm.get("include_disclaimer"):
        disclaimer_html = f"<p><strong>Disclaimer:</strong> {fm.get('disclaimer_text','')}</p>"
    publisher_html = f"<p>Publisher: {fm.get('publisher_name')}</p>" if fm.get("publisher_name") else ""
    isbn_html = f"<p>ISBN: {fm.get('isbn')}</p>" if fm.get("isbn") else ""
    copyright_page.content = f"""
    <html>
    <head><title>Copyright</title></head>
    <body>
        <h1>Copyright</h1>
        <p>© {fm['copyright_year']} {fm['author_name']}. {fm['rights_statement']}</p>
        {publisher_html}
        {isbn_html}
        {disclaimer_html}
    </body>
    </html>
    """
    book.add_item(copyright_page)
    epub_chapters.append(copyright_page)

    # Also By (optional)
    if sup["also_by"]:
        also_by_page = epub.EpubHtml(title="Also By", file_name="also_by.xhtml", lang="en")
        items = "".join(f"<li>{t}</li>" for t in sup["also_by"])
        also_by_page.content = f"""
        <html>
        <head><title>Also By</title></head>
        <body>
            <h1>Also By</h1>
            <ul>{items}</ul>
        </body>
        </html>
        """
        book.add_item(also_by_page)
        epub_chapters.append(also_by_page)

    if chapters:
        for chapter in sorted(chapters, key=lambda x: x.get('number', 0)):
            ch_num = chapter.get('number', '?')
            ch_title = chapter.get('title', f'Chapter {ch_num}')

            # Create chapter
            ch = epub.EpubHtml(
                title=f'Chapter {ch_num}: {ch_title}',
                file_name=f'chapter_{ch_num}.xhtml',
                lang='en'
            )

            # Format chapter content
            text = chapter.get('text', '')
            if text:
                # Convert text to HTML paragraphs
                paragraphs = text.split('\n\n')
                html_content = f'<h1>Chapter {ch_num}: {ch_title}</h1>\n'

                for para in paragraphs:
                    para = para.strip()
                    if para:
                        # Handle scene breaks
                        if para in ['* * *', '---', '***']:
                            html_content += '<p style="text-align: center; margin: 2em 0;">* * *</p>\n'
                        else:
                            # Escape HTML entities
                            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            html_content += f'<p style="text-indent: 1.5em; margin: 0.5em 0;">{para}</p>\n'
            else:
                html_content = f'''
                <h1>Chapter {ch_num}: {ch_title}</h1>
                <p><em>Chapter content not yet written.</em></p>
                '''

            ch.content = f'''
            <html>
            <head><title>Chapter {ch_num}</title></head>
            <body>
                {html_content}
            </body>
            </html>
            '''

            book.add_item(ch)
            epub_chapters.append(ch)
    else:
        # No chapters yet
        empty_ch = epub.EpubHtml(title='No Content', file_name='empty.xhtml', lang='en')
        empty_ch.content = '''
        <html>
        <head><title>No Content</title></head>
        <body>
            <h1>No Chapters Written</h1>
            <p>Use the Chapter Writer to generate content from your outline.</p>
        </body>
        </html>
        '''
        book.add_item(empty_ch)
        epub_chapters.append(empty_ch)

    # Back matter (optional)
    if sup.get("acknowledgements"):
        acks = epub.EpubHtml(title="Acknowledgements", file_name="acknowledgements.xhtml", lang="en")
        txt = sup["acknowledgements"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        acks.content = f"<html><head><title>Acknowledgements</title></head><body><h1>Acknowledgements</h1><p>{txt}</p></body></html>"
        book.add_item(acks)
        epub_chapters.append(acks)

    if sup.get("newsletter_cta") or sup.get("newsletter_url"):
        news = epub.EpubHtml(title="Stay in Touch", file_name="newsletter.xhtml", lang="en")
        cta = sup.get("newsletter_cta", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        url = sup.get("newsletter_url", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        parts = []
        if cta:
            parts.append(f"<p>{cta}</p>")
        if url:
            parts.append(f"<p>{url}</p>")
        news.content = f"<html><head><title>Stay in Touch</title></head><body><h1>Stay in Touch</h1>{''.join(parts)}</body></html>"
        book.add_item(news)
        epub_chapters.append(news)

    if sup.get("about_author"):
        about = epub.EpubHtml(title="About the Author", file_name="about_author.xhtml", lang="en")
        txt = sup["about_author"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        about.content = f"<html><head><title>About the Author</title></head><body><h1>About the Author</h1><p>{txt}</p></body></html>"
        book.add_item(about)
        epub_chapters.append(about)

    # Define Table of Contents
    book.toc = tuple(epub_chapters)

    # Add navigation files
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # Define CSS style
    style = '''
    body {
        font-family: Georgia, serif;
        margin: 1em;
        line-height: 1.6;
    }
    h1 {
        text-align: center;
        margin-bottom: 1.5em;
        font-size: 1.5em;
    }
    p {
        text-indent: 1.5em;
        margin: 0.5em 0;
    }
    '''
    nav_css = epub.EpubItem(
        uid="style_nav",
        file_name="style/nav.css",
        media_type="text/css",
        content=style
    )
    book.add_item(nav_css)

    # Create spine
    book.spine = ['nav'] + epub_chapters

    # Write to bytes
    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    buffer.seek(0)
    return buffer.getvalue()


def generate_kindle_mobi(project) -> Optional[bytes]:
    """
    Generate a MOBI file for Kindle.
    Note: This requires kindlegen or calibre's ebook-convert to be installed.
    Returns None if conversion tools aren't available.

    For now, we'll just return the EPUB since modern Kindles support EPUB,
    and users can use Calibre to convert if needed.
    """
    # MOBI generation typically requires external tools
    # For simplicity, we recommend using the EPUB and converting with Calibre
    # or Amazon's Kindle Previewer
    return None


def get_word_count(project) -> int:
    """Calculate total word count of written chapters."""
    chapters = project.manuscript.get('chapters', [])
    return sum(ch.get('word_count', 0) for ch in chapters)


def get_chapter_summary(project) -> List[Dict[str, Any]]:
    """Get summary of all chapters."""
    chapters = project.manuscript.get('chapters', [])
    return [
        {
            'number': ch.get('number'),
            'title': ch.get('title'),
            'word_count': ch.get('word_count', 0),
            'written': bool(ch.get('text'))
        }
        for ch in sorted(chapters, key=lambda x: x.get('number', 0))
    ]
